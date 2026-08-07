#!/usr/bin/env python3
"""Create a test .docx template for handbook-generator skill testing."""

import sys
sys.path.insert(0, r'C:\Users\Administrator\.codebuddy\skills\handbook-generator\scripts')

# Auto-install dependencies
from utils import ensure_dependencies
ensure_dependencies()

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Title ----
title = doc.add_heading('{{产品名称}} 产品手册', level=0)

# ---- Basic Info ----
doc.add_heading('基本信息', level=1)
table = doc.add_table(rows=4, cols=2, style='Table Grid')
table.cell(0, 0).text = '产品名称'
table.cell(0, 1).text = '{{产品名称}}'
table.cell(1, 0).text = '适用人群'
table.cell(1, 1).text = '{{适用人群}}'
table.cell(2, 0).text = '保障期限'
table.cell(2, 1).text = '{{保障期限}}'
table.cell(3, 0).text = '最低保费'
table.cell(3, 1).text = '{{最低保费}}元/年'

# ---- Product Image ----
doc.add_heading('产品示意图', level=1)
doc.add_paragraph('{{img_产品示意图}}')

# ---- Coverage Details ----
doc.add_heading('保障详情', level=1)
doc.add_paragraph('{{保障详情描述}}')

# Coverage table
table2 = doc.add_table(rows=3, cols=3, style='Table Grid')
table2.cell(0, 0).text = '保障项目'
table2.cell(0, 1).text = '保额'
table2.cell(0, 2).text = '说明'
table2.cell(1, 0).text = '{{保障项目1}}'
table2.cell(1, 1).text = '{{保额1}}'
table2.cell(1, 2).text = '{{保障说明1}}'
table2.cell(2, 0).text = '{{保障项目2}}'
table2.cell(2, 1).text = '{{保额2}}'
table2.cell(2, 2).text = '{{保障说明2}}'

# ---- Claims Process ----
doc.add_heading('理赔流程', level=1)
doc.add_paragraph('{{理赔流程说明}}')

# ---- Company Logo ----
doc.add_heading('公司信息', level=1)
doc.add_paragraph('{{img_公司Logo}}')
doc.add_paragraph('客服电话：{{客服电话}}')
doc.add_paragraph('官方网站：{{官方网站}}')

# Save
output = r'C:\Users\Administrator\.codebuddy\skills\handbook-generator\test_template.docx'
doc.save(output)
print(f'Test template created: {output}')
