#!/usr/bin/env python3
"""
Handbook Generator - Modify Mode

Modifies an existing .docx handbook by:
  a. Re-filling {{placeholder}} and {{img_name}} placeholders (re-generation)
  b. Replacing specific paragraph/text content by search-and-replace

Usage:
    modify.py <docx_path> [options]

Options:
    --fill-data <json>     JSON string with placeholder values (text + img paths)
    --replace <json>       JSON string with {old_text: new_text} pairs for paragraph replacement
    --output <path>        Output file path (default: overwrite original)
    --open                 Open the modified file after saving
    --scan                 Scan and display remaining placeholders in the document
"""

import sys
import json
from pathlib import Path

# Ensure utils is importable from the same directory
sys.path.insert(0, str(Path(__file__).parent))
from utils import (
    ensure_dependencies,
    scan_placeholders,
    format_placeholders_list,
    open_file,
    validate_template_path,
    validate_image_path,
)


def modify_by_replacement(docx_path, replacements, output_path, should_open=False):
    """
    Replace specific text content in an existing .docx file.

    Args:
        docx_path: Path to the .docx file to modify
        replacements: Dict mapping old_text strings to new_text strings
        output_path: Path for the modified output file
        should_open: Whether to open the file after modification
    """
    from docx import Document

    docx_path = validate_template_path(docx_path)
    doc = Document(str(docx_path))

    replaced_count = 0

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                # Need to replace at run level to preserve formatting
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replaced_count += 1

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replaced_count += 1

    # Save
    if output_path is None:
        output_path = str(docx_path)
    doc.save(output_path)
    print(f"Replaced {replaced_count} occurrences. Saved to: {output_path}")

    if should_open:
        open_file(output_path)


def modify_by_refill(docx_path, text_data, img_data, output_path, should_open=False):
    """
    Re-fill placeholders in an existing .docx file (same as generate, but for modification).

    Args:
        docx_path: Path to the .docx template file
        text_data: Dict mapping text placeholder names to their values
        img_data: Dict mapping image placeholder names to image file paths
        output_path: Path for the modified output file
        should_open: Whether to open the file after modification
    """
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm

    docx_path = validate_template_path(docx_path)

    # Scan template to discover all placeholders (for auto-filling missing ones)
    placeholders = scan_placeholders(str(docx_path))

    # Build the context dict
    context = {}
    for name in placeholders['text']:
        if name in text_data:
            context[name] = text_data[name]
        else:
            context[name] = f"[待填写: {name}]"

    tpl = DocxTemplate(str(docx_path))
    for img_name in placeholders['img']:
        context_key = f"img_{img_name}"
        if img_name in img_data:
            img_path = validate_image_path(img_data[img_name])
            if img_path:
                context[context_key] = InlineImage(tpl, str(img_path), width=Mm(150))
            else:
                context[context_key] = f"[待插入图片: {img_name}]"
        else:
            context[context_key] = f"[待插入图片: {img_name}]"

    tpl.render(context)

    if output_path is None:
        output_path = str(docx_path)
    tpl.save(str(output_path))
    print(f"Handbook re-filled successfully: {output_path}")

    if should_open:
        open_file(str(output_path))


def main():
    import argparse

    parser = argparse.ArgumentParser(description='Modify an existing .docx handbook')
    parser.add_argument('docx', help='Path to the .docx file to modify')
    parser.add_argument('--fill-data', help='JSON string with placeholder values for re-filling', default=None)
    parser.add_argument('--replace', help='JSON string with {old: new} text replacements', default=None)
    parser.add_argument('--output', help='Output file path (default: overwrite original)', default=None)
    parser.add_argument('--open', action='store_true', help='Open the modified file after saving')
    parser.add_argument('--scan', action='store_true', help='Scan and display remaining placeholders')

    args = parser.parse_args()

    # Ensure dependencies are available
    ensure_dependencies()

    # Scan mode
    if args.scan:
        placeholders = scan_placeholders(args.docx)
        print(format_placeholders_list(placeholders))
        return

    # Fill mode: re-fill placeholders
    if args.fill_data:
        try:
            fill_data = json.loads(args.fill_data)
        except json.JSONDecodeError as e:
            print(f"Error parsing fill-data JSON: {e}")
            sys.exit(1)

        # Separate text and image data
        text_data = {}
        img_data = {}
        for key, value in fill_data.items():
            # Image keys have "img_" prefix in the data (Jinja2-compatible)
            if key.startswith('img_'):
                img_data[key[4:]] = value
            else:
                text_data[key] = value

        modify_by_refill(
            docx_path=args.docx,
            text_data=text_data,
            img_data=img_data,
            output_path=args.output,
            should_open=args.open
        )

    # Replace mode: text search-and-replace
    if args.replace:
        try:
            replacements = json.loads(args.replace)
        except json.JSONDecodeError as e:
            print(f"Error parsing replace JSON: {e}")
            sys.exit(1)

        modify_by_replacement(
            docx_path=args.docx,
            replacements=replacements,
            output_path=args.output,
            should_open=args.open
        )

    if not args.fill_data and not args.replace and not args.scan:
        print("Error: Must specify at least one modification mode (--fill-data, --replace, or --scan)")
        sys.exit(1)


if __name__ == '__main__':
    main()
