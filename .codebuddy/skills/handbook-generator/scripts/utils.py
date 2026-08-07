#!/usr/bin/env python3
"""
Handbook Generator - Shared Utilities

Provides common functions for template scanning, dependency checking,
and document operations used by generate.py and modify.py.
"""

import sys
import os
import re
import subprocess
from pathlib import Path


def ensure_dependencies():
    """Check and auto-install required Python packages if missing."""
    required = {
        'docx': 'python-docx',
        'docxtpl': 'docxtpl',
    }
    missing = []
    for module, package in required.items():
        try:
            __import__(module)
        except ImportError:
            missing.append(package)

    if missing:
        print(f"Installing missing dependencies: {', '.join(missing)}")
        subprocess.check_call(
            [sys.executable, '-m', 'pip', 'install', '--quiet'] + missing
        )
        print("Dependencies installed successfully.")


def scan_placeholders(docx_path):
    """
    Scan a .docx template file for all {{variable}} and {{img_name}} placeholders.

    Note: Image placeholders use {{img_name}} syntax (underscore) instead of {{img:name}}
    because Jinja2 does not allow colons in variable names.

    Returns:
        dict with keys:
            'text': list of text placeholder names (e.g. ['产品名称', '投保流程'])
            'img': list of image placeholder names (e.g. ['产品示意图', '公司Logo'])
                  These are the names AFTER removing the 'img_' prefix.
    """
    from docx import Document

    doc = Document(docx_path)
    text_placeholders = []
    img_placeholders = []

    # Pattern for all placeholders: {{name}}
    text_pattern = re.compile(r'\{\{([^{}]+)\}\}')
    # Pattern for image placeholders: {{img_name}} (underscore, Jinja2-compatible)
    img_pattern = re.compile(r'img_(.+)')

    # Scan all paragraphs
    for paragraph in doc.paragraphs:
        for match in text_pattern.finditer(paragraph.text):
            name = match.group(1).strip()
            img_match = img_pattern.match(name)
            if img_match:
                img_placeholders.append(img_match.group(1).strip())
            else:
                text_placeholders.append(name)

    # Scan all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for match in text_pattern.finditer(paragraph.text):
                        name = match.group(1).strip()
                        img_match = img_pattern.match(name)
                        if img_match:
                            img_placeholders.append(img_match.group(1).strip())
                        else:
                            text_placeholders.append(name)

    # Deduplicate while preserving order
    text_placeholders = list(dict.fromkeys(text_placeholders))
    img_placeholders = list(dict.fromkeys(img_placeholders))

    return {'text': text_placeholders, 'img': img_placeholders}


def format_placeholders_list(placeholders):
    """Format scanned placeholders into a readable display string."""
    lines = []
    if placeholders['text']:
        lines.append("=== Text Placeholders ===")
        for i, name in enumerate(placeholders['text'], 1):
            lines.append(f"  {i}. {{{{{name}}}}}")
    if placeholders['img']:
        lines.append("=== Image Placeholders ===")
        for i, name in enumerate(placeholders['img'], 1):
            lines.append(f"  {i}. {{{{img_{name}}}}}")
    if not lines:
        lines.append("No placeholders found in the template.")
    return '\n'.join(lines)


def resolve_output_path(template_path, output_path=None):
    """
    Resolve the output file path for the generated handbook.

    If output_path is provided by the user, use it directly.
    Otherwise, save to the same directory as the template with a _generated suffix.
    """
    if output_path:
        return Path(output_path).resolve()

    template = Path(template_path).resolve()
    return template.parent / f"{template.stem}_generated{template.suffix}"


def open_file(filepath):
    """Open the generated file using the system's default application."""
    filepath = str(Path(filepath).resolve())
    if sys.platform == 'win32':
        os.startfile(filepath)
    elif sys.platform == 'darwin':
        subprocess.run(['open', filepath])
    else:
        subprocess.run(['xdg-open', filepath])


def validate_template_path(docx_path):
    """Validate that the template file exists and is a .docx file."""
    path = Path(docx_path).resolve()
    if not path.exists():
        print(f"Error: Template file not found: {path}")
        sys.exit(1)
    if path.suffix.lower() != '.docx':
        print(f"Error: Template file must be a .docx file, got: {path.suffix}")
        sys.exit(1)
    return path


def validate_image_path(img_path):
    """Validate that an image file exists."""
    path = Path(img_path).resolve()
    if not path.exists():
        print(f"Warning: Image file not found: {path}")
        return None
    return path
